from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict
import logging
import os

class WordDocumentGenerator:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def create_document_from_chunks(self, translated_chunks: List[Dict], output_path: str, 
                                   title: str = "Translated Document") -> str:
        """Create a Word document from translated chunks"""
        try:
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add document info
            doc.add_paragraph(f"Translation completed with {len(translated_chunks)} chunks")
            doc.add_paragraph("")  # Empty line
            
            # Add content from each chunk
            for chunk_idx, chunk in enumerate(translated_chunks, 1):
                # Add chunk header
                chunk_header = f"Pages {chunk['start_page']}"
                if chunk['start_page'] != chunk['end_page']:
                    chunk_header += f"-{chunk['end_page']}"
                
                doc.add_heading(chunk_header, level=2)
                
                # Add pages content
                for page_num, translated_text in chunk['pages']:
                    if translated_text.strip():
                        # Add page number as subtitle
                        page_header = doc.add_heading(f"Page {page_num}", level=3)
                        
                        # Add translated content
                        paragraphs = translated_text.split('\n')
                        for paragraph_text in paragraphs:
                            if paragraph_text.strip():
                                doc.add_paragraph(paragraph_text.strip())
                        
                        # Add spacing between pages
                        doc.add_paragraph("")
                
                # Add page break between chunks (except for the last one)
                if chunk_idx < len(translated_chunks):
                    doc.add_page_break()
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"Word document saved to: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error creating Word document: {str(e)}")
            raise
    
    def create_bilingual_document(self, original_chunks: List[Dict], translated_chunks: List[Dict], 
                                 output_path: str, title: str = "Bilingual Document") -> str:
        """Create a bilingual Word document with original and translated text side by side"""
        try:
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add document info
            doc.add_paragraph("Bilingual Translation Document")
            doc.add_paragraph("Left: Original Chinese | Right: English Translation")
            doc.add_paragraph("")
            
            # Process chunks
            for chunk_idx, (orig_chunk, trans_chunk) in enumerate(zip(original_chunks, translated_chunks), 1):
                # Add chunk header
                chunk_header = f"Pages {orig_chunk['start_page']}"
                if orig_chunk['start_page'] != orig_chunk['end_page']:
                    chunk_header += f"-{orig_chunk['end_page']}"
                
                doc.add_heading(chunk_header, level=2)
                
                # Create table for bilingual content
                for (orig_page_num, orig_text), (trans_page_num, trans_text) in zip(orig_chunk['pages'], trans_chunk['pages']):
                    if orig_text.strip() or trans_text.strip():
                        # Add page number
                        doc.add_heading(f"Page {orig_page_num}", level=3)
                        
                        # Create table for bilingual content
                        table = doc.add_table(rows=1, cols=2)
                        table.style = 'Table Grid'
                        
                        # Set column widths
                        table.columns[0].width = Inches(3)
                        table.columns[1].width = Inches(3)
                        
                        # Add headers
                        header_cells = table.rows[0].cells
                        header_cells[0].text = "Original (Chinese)"
                        header_cells[1].text = "Translation (English)"
                        
                        # Add content
                        row_cells = table.add_row().cells
                        row_cells[0].text = orig_text
                        row_cells[1].text = trans_text
                        
                        doc.add_paragraph("")  # Spacing
                
                # Add page break between chunks
                if chunk_idx < len(original_chunks):
                    doc.add_page_break()
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"Bilingual Word document saved to: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error creating bilingual Word document: {str(e)}")
            raise
    
    def create_simple_document(self, pages_text: List[tuple], output_path: str, 
                              title: str = "Translated Document") -> str:
        """Create a simple Word document from a list of (page_num, text) tuples"""
        try:
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
            
            # Add content
            for page_num, text in pages_text:
                if text.strip():
                    # Add page header
                    doc.add_heading(f"Page {page_num}", level=2)
                    
                    # Add content paragraphs
                    paragraphs = text.split('\n')
                    for paragraph_text in paragraphs:
                        if paragraph_text.strip():
                            doc.add_paragraph(paragraph_text.strip())
                    
                    doc.add_paragraph("")  # Add spacing
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"Simple Word document saved to: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error creating simple Word document: {str(e)}")
            raise
    
    def generate_output_filename(self, input_pdf_path: str, suffix: str = "_translated") -> str:
        """Generate output filename based on input PDF path"""
        base_name = os.path.splitext(os.path.basename(input_pdf_path))[0]
        output_dir = os.path.dirname(input_pdf_path)
        return os.path.join(output_dir, f"{base_name}{suffix}.docx")